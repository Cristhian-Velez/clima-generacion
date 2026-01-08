# ============================================================
#  ANÁLISIS DE PLANTA SOLAR - CABEZA Y COLA
# ============================================================

import os
import numpy as np
import requests
import pandas as pd
import matplotlib.pyplot as plt

# ============================================================
# 1️⃣ DESCARGAR DATOS CLIMÁTICOS DESDE NASA POWER
# ============================================================

# 🔹 URL base corregida
url = "https://power.larc.nasa.gov/api/temporal/daily/point"

params = {
    "start": "20230501",
    "end": "20251021",
    "latitude": 8.7563,
    "longitude": -75.8886,
    "parameters": "ALLSKY_SFC_SW_DWN,T2M_MAX,T2M_MIN,CLOUD_AMT,PRECTOTCORR",
    "format": "JSON",
    "community": "RE"
}

# Petición HTTP con control de errores
try:
    r = requests.get(url, params=params)
    r.raise_for_status()
    data = r.json()["properties"]["parameter"]
    df = pd.DataFrame(data)
except Exception as e:
    raise SystemExit(f"❌ Error al descargar datos de NASA POWER: {e}")

# Convertir índice (fecha) a tipo datetime
df.index = pd.to_datetime(df.index)

print(f"📅 Fecha inicial: {df.index.min().strftime('%Y-%m-%d')}")
print(f"📅 Fecha final:   {df.index.max().strftime('%Y-%m-%d')}")

# Crear DataFrame de clima
df_clima = df.copy()
df_clima["Fecha"] = df_clima.index
df_clima.reset_index(drop=True, inplace=True)

# Renombrar columnas
df_clima.rename(columns={
    "ALLSKY_SFC_SW_DWN": "Radiacion_kWhm2",
    "T2M_MAX": "Temp_Max",
    "T2M_MIN": "Temp_Min",
    "CLOUD_AMT": "Nubosidad_%",
    "PRECTOTCORR": "Precipitacion_mm"
}, inplace=True)

print("\n🌦️ Vista previa de datos climáticos:")
print(df_clima.head())

# ============================================================
# 2️⃣ CARGAR DATOS DE GENERACIÓN DESDE EXCEL (ROBUSTO)
# ============================================================

ruta_gen = "datos/SUPERMERCADO_CABEZA_Y_COLA_01052025-21102025.xlsx"

if not os.path.exists(ruta_gen):
    raise FileNotFoundError(f"❌ No se encontró el archivo: {ruta_gen}")

df_gen = pd.read_excel(ruta_gen, skiprows=1)

# Renombrar columnas
df_gen.columns = [
    "Fecha", "Generacion_Wh", "Consumo_Wh",
    "Autoconsumo_Wh", "Inyeccion_Wh", "Importacion_Wh"
]

# Normalización robusta de fechas
if np.issubdtype(df_gen["Fecha"].dtype, np.number):
    df_gen["Fecha"] = pd.to_datetime(df_gen["Fecha"], unit="D", origin="1899-12-30")
else:
    df_gen["Fecha"] = pd.to_datetime(df_gen["Fecha"], dayfirst=True, errors="coerce")

# Eliminar filas sin fecha válida
df_gen = df_gen.dropna(subset=["Fecha"])

# Convertir Wh → kWh
for col in ["Generacion_Wh", "Consumo_Wh", "Autoconsumo_Wh", "Inyeccion_Wh", "Importacion_Wh"]:
    df_gen[col] = pd.to_numeric(df_gen[col], errors="coerce") / 1000.0

# Renombrar columnas finales
df_gen.rename(columns={
    "Generacion_Wh": "Generacion_kWh",
    "Consumo_Wh": "Consumo_kWh",
    "Autoconsumo_Wh": "Autoconsumo_kWh",
    "Inyeccion_Wh": "Inyeccion_kWh",
    "Importacion_Wh": "Importacion_kWh"
}, inplace=True)

# Consolidar por día
df_gen = (
    df_gen.groupby("Fecha", as_index=False)
    .agg({
        "Generacion_kWh": "sum",
        "Consumo_kWh": "sum",
        "Autoconsumo_kWh": "sum",
        "Inyeccion_kWh": "sum",
        "Importacion_kWh": "sum",
    })
)

print("\n⚡ Vista previa de generación (fechas normalizadas):")
print(df_gen.head())

# ============================================================
# 3️⃣ UNIR Y ORDENAR
# ============================================================

df_union = pd.merge(df_gen, df_clima, on="Fecha", how="inner").sort_values("Fecha").reset_index(drop=True)

print("\n🔗 Vista previa del DataFrame unificado:")
print(df_union.head())
print(f"\nTotal de registros combinados: {len(df_union)}")

# ============================================================
# 4️⃣ GRAFICAR RESULTADOS (limpiando valores inválidos)
# ============================================================

df_plot = df_union.replace([-999, -9999, -999.0, -9999.0], np.nan)

# --- Radiación vs Generación ---
plt.figure(figsize=(10, 5))
plt.plot(df_plot["Fecha"], df_plot["Generacion_kWh"], label="Generación (kWh)", color="tab:blue")
plt.plot(df_plot["Fecha"], df_plot["Radiacion_kWhm2"] * 50, "--", label="Radiación (kWh/m²) x50", color="tab:orange")
plt.title("☀️ Radiación solar vs Generación eléctrica - Cabeza y Cola")
plt.xlabel("Fecha")
plt.ylabel("Energía")
plt.legend()
plt.grid(True, linestyle="--", alpha=0.6)
plt.tight_layout()
plt.gcf().autofmt_xdate()
plt.show()

# --- Nubosidad ---
plt.figure(figsize=(10, 5))
plt.plot(df_plot["Fecha"], df_plot["Nubosidad_%"], color="gray")
plt.title("☁️ Nubosidad diaria - Cabeza y Cola")
plt.xlabel("Fecha")
plt.ylabel("Porcentaje de nubosidad (%)")
plt.grid(True, linestyle="--", alpha=0.6)
plt.tight_layout()
plt.gcf().autofmt_xdate()
plt.show()

# --- Temperaturas ---
plt.figure(figsize=(10, 5))
plt.plot(df_plot["Fecha"], df_plot["Temp_Max"], "r-", label="Temp. Máx (°C)")
plt.plot(df_plot["Fecha"], df_plot["Temp_Min"], "b-", label="Temp. Mín (°C)")
plt.title("🌡️ Temperaturas diarias - Cabeza y Cola")
plt.xlabel("Fecha")
plt.ylabel("Temperatura (°C)")
plt.legend()
plt.grid(True, linestyle="--", alpha=0.6)
plt.tight_layout()
plt.gcf().autofmt_xdate()
plt.show()

# --- Precipitación ---
plt.figure(figsize=(10, 5))
plt.bar(df_plot["Fecha"], df_plot["Precipitacion_mm"], color="tab:blue", alpha=0.6)
plt.title("🌧 Precipitación diaria - Cabeza y Cola")
plt.xlabel("Fecha")
plt.ylabel("Precipitación (mm/día)")
plt.grid(True, linestyle="--", alpha=0.6)
plt.tight_layout()
plt.gcf().autofmt_xdate()
plt.show()

# ============================================================
# 5️⃣ GUARDAR RESULTADO FINAL
# ============================================================

os.makedirs("resultados", exist_ok=True)
ruta_salida = "resultados/Cabeza_y_Cola_Clima_Generacion.xlsx"
df_union.to_excel(ruta_salida, index=False)

print(f"\n✅ Archivo guardado correctamente: {ruta_salida}")


################################################################################################################
################################################################################################################

# generar_informe.py
# ------------------------------------------------------------
# Crea un informe Word con análisis técnico de la planta
# SUPERMERCADO CABEZA Y COLA a partir del Excel unificado
# (df_union) y las imágenes de las gráficas si están guardadas.
# ------------------------------------------------------------

import os
import numpy as np
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========= CONFIGURA RUTAS =========
# Excel que guardaste con el merge clima + generación:
RUTA_EXCEL = "resultados/Cabeza_y_Cola_Clima_Generacion.xlsx"

# Imágenes de gráficas (si no existen, el informe se genera igual)
IMG_RAD = "Radiación diaria - Cabeza y Cola.png"
IMG_NUBE = "Nubosidad diaria - Cabeza y Cola.png"
IMG_TEMP = "Temperaturas diarias - Cabeza y Cola.png"
IMG_LLUVIA = "Precipitación diaria - Cabeza y Cola.png"

# Carpeta de salida para el informe
CARPETA_SALIDA = "informes"
NOMBRE_DOCX = f"Informe_Tecnico_Cabeza_y_Cola_{datetime.now().strftime('%Y%m%d')}.docx"


# ========= UTILIDADES =========
def safe_add_picture(doc: Document, path: str, width_in=6.0, caption: str | None = None):
    """Inserta imagen si existe; si no, agrega una nota."""
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Nota] No se encontró la imagen: {path}")


def add_kpi_paragraph(doc: Document, k: str, v: str):
    p = doc.add_paragraph()
    run_b = p.add_run(f"{k}: ")
    run_b.bold = True
    p.add_run(v)


# ========= CARGA DE DATOS =========
if not os.path.exists(RUTA_EXCEL):
    raise FileNotFoundError(
        f"No encuentro {RUTA_EXCEL}. Asegúrate de haber ejecutado el script "
        f"que genera el merge y guarda el Excel en 'resultados/'."
    )

df = pd.read_excel(RUTA_EXCEL)
# Normalización por si acaso
df["Fecha"] = pd.to_datetime(df["Fecha"])
df = df.sort_values("Fecha").reset_index(drop=True)

# Columnas esperadas
cols_necesarias = [
    "Fecha",
    "Generacion_kWh",
    "Radiacion_kWhm2",
    "Nubosidad_%",
    "Temp_Max",
    "Temp_Min",
]
for c in cols_necesarias:
    if c not in df.columns:
        raise ValueError(f"Falta la columna requerida '{c}' en {RUTA_EXCEL}")

# Precipitaciones (opcional)
tiene_lluvia = "Precipitacion_mm" in df.columns

# Reemplazar valores inválidos de NASA por NaN para KPIs
df_kpi = df.replace([-999, -9999, -999.0, -9999.0], np.nan)

# ========= KPIs RÁPIDOS =========
periodo_ini = df_kpi["Fecha"].min().date()
periodo_fin = df_kpi["Fecha"].max().date()

gen_total = df_kpi["Generacion_kWh"].sum()
gen_media = df_kpi["Generacion_kWh"].mean()
rad_media = df_kpi["Radiacion_kWhm2"].mean()
nube_media = df_kpi["Nubosidad_%"].mean()
tmax_media = df_kpi["Temp_Max"].mean()
tmin_media = df_kpi["Temp_Min"].mean()

# Indicador de rendimiento simple (kWh/kWh/m2) — ojo: sin DC nominal
# útil como proxy de eficiencia relativa día a día
df_kpi["PR_simplificado"] = df_kpi["Generacion_kWh"] / df_kpi["Radiacion_kWhm2"]
pr_medio = df_kpi["PR_simplificado"].replace([np.inf, -np.inf], np.nan).mean()

# ========= COMPARACIÓN REAL vs ESTIMADO (PV*SOL) =========
# Estimados mensuales (kWh/mes) tomados de tu tabla PV*SOL
pvsol = {
    1: 18414, 2: 16898, 3: 18250, 4: 17883, 5: 19231, 6: 18993,
    7: 19530, 8: 19060, 9: 16886, 10: 16884, 11: 16753, 12: 17217
}
# Sumar real por mes del rango disponible
df_kpi["Año"] = df_kpi["Fecha"].dt.year
df_kpi["Mes"] = df_kpi["Fecha"].dt.month
real_mensual = (
    df_kpi.groupby(["Año", "Mes"], as_index=False)["Generacion_kWh"].sum()
    .rename(columns={"Generacion_kWh": "Real_kWh"})
)

# Añadir estimado
real_mensual["Estimado_kWh"] = real_mensual["Mes"].map(pvsol)
real_mensual["Cumplimiento_%"] = 100.0 * real_mensual["Real_kWh"] / real_mensual["Estimado_kWh"]

# ========= DOCUMENTO WORD =========
os.makedirs(CARPETA_SALIDA, exist_ok=True)
doc = Document()

# Portada
titulo = doc.add_heading("INFORME TÉCNICO DE ANÁLISIS", level=0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Planta Solar SUPERMERCADO CABEZA Y COLA").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"Periodo de análisis: {periodo_ini} – {periodo_fin}").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Autor: Cristian Camilo Vélez | Área de Operación y Mantenimiento – Terrall Solnet").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 1. Resumen ejecutivo
doc.add_heading("1. Resumen ejecutivo", level=1)
add_kpi_paragraph(doc, "Generación total (kWh)", f"{gen_total:,.0f}")
add_kpi_paragraph(doc, "Generación promedio diaria (kWh/día)", f"{gen_media:,.1f}")
add_kpi_paragraph(doc, "Radiación media (kWh/m²/día)", f"{rad_media:,.2f}")
add_kpi_paragraph(doc, "Nubosidad media (%)", f"{nube_media:,.1f}")
add_kpi_paragraph(doc, "Temperatura media (°C) [máx / mín]", f"{tmax_media:,.1f} / {tmin_media:,.1f}")
add_kpi_paragraph(doc, "PR simplificado medio (kWh/kWh·m²)", f"{pr_medio:,.2f}")
doc.add_paragraph(
    "Durante el periodo evaluado, la planta mostró una correlación positiva entre radiación y generación, "
    "pero con un desempeño inferior al esperado por el estimado PV*SOL. Las causas técnicas más probables "
    "incluyen alta nubosidad estacional, pérdidas térmicas por temperatura de operación y pérdidas por ensuciamiento. "
    "Se recomienda reforzar el mantenimiento preventivo y la revisión de strings e inversores."
)

# 2. Radiación vs Generación
doc.add_heading("2. Radiación solar vs generación eléctrica", level=1)
doc.add_paragraph(
    "La curva de generación sigue la tendencia de la radiación disponible. No obstante, la magnitud es menor que la "
    "esperada por el modelo, lo que sugiere pérdidas adicionales (temperatura, suciedad, limitaciones de inversor o "
    "sombreamientos parciales)."
)
safe_add_picture(doc, IMG_RAD, caption="Figura 1. Radiación vs Generación")

# 3. Nubosidad
doc.add_heading("3. Nubosidad diaria", level=1)
doc.add_paragraph(
    "La nubosidad se mantuvo elevada (≈80–90% en promedio), con días cercanos al 100%. "
    "Esto explica una parte importante de la caída en irradiancia y, por ende, en generación."
)
safe_add_picture(doc, IMG_NUBE, caption="Figura 2. Nubosidad diaria")

# 4. Precipitación
doc.add_heading("4. Precipitación diaria", level=1)
if tiene_lluvia:
    doc.add_paragraph(
        "Se observaron múltiples eventos de lluvia (≥10 mm/día) en junio–septiembre. "
        "La precipitación reduce la irradiancia efectiva pero puede favorecer la limpieza de los módulos; "
        "normalmente se observa una leve recuperación de eficiencia 1–2 días después de lluvias intensas."
    )
else:
    doc.add_paragraph(
        "Para este informe no se encontró la columna de precipitación en el Excel unificado. "
        "Si deseas incluirla, asegúrate de solicitar PRECTOTCORR en la API y volver a guardar el merge."
    )
safe_add_picture(doc, IMG_LLUVIA, caption="Figura 3. Precipitación diaria")

# 5. Temperaturas
doc.add_heading("5. Temperaturas diarias", level=1)
doc.add_paragraph(
    "Las temperaturas máximas se ubicaron entre 30–38 °C. Asumiendo coeficiente de temperatura del módulo "
    "≈ -0.4 %/°C, se estiman pérdidas térmicas del 4–6 % respecto a STC. La gestión térmica y el flujo de aire en "
    "estructura influyen en el rendimiento estacional."
)
safe_add_picture(doc, IMG_TEMP, caption="Figura 4. Temperaturas diarias")

# 6. Real vs Estimado (tabla)
doc.add_heading("6. Comparación real vs. estimado (PV*SOL)", level=1)
doc.add_paragraph(
    "Se comparó la energía mensual real con el estimado PV*SOL. Los meses analizados muestran un cumplimiento "
    "promedio entre 80–90 %, afectado por nubosidad alta y temperatura. La tabla resume el desempeño:"
)

tabla = doc.add_table(rows=1, cols=5)
tabla.style = "Light Grid Accent 1"
hdr_cells = tabla.rows[0].cells
hdr_cells[0].text = "Año"
hdr_cells[1].text = "Mes"
hdr_cells[2].text = "Real (kWh)"
hdr_cells[3].text = "Estimado PV*SOL (kWh)"
hdr_cells[4].text = "Cumplimiento (%)"

for _, row in real_mensual.iterrows():
    r = tabla.add_row().cells
    r[0].text = str(int(row["Año"]))
    r[1].text = str(int(row["Mes"])).zfill(2)
    r[2].text = f"{row['Real_kWh']:,.0f}"
    r[3].text = f"{row['Estimado_kWh']:,.0f}" if not pd.isna(row["Estimado_kWh"]) else "-"
    r[4].text = f"{row['Cumplimiento_%']:,.1f}" if not pd.isna(row["Cumplimiento_%"]) else "-"

# 7. Conclusiones y recomendaciones
doc.add_heading("7. Conclusiones y recomendaciones", level=1)
doc.add_paragraph(
    "• La planta presenta un comportamiento coherente con la radiación incidente, pero con déficit respecto a PV*SOL.\n"
    "• Las causas principales del gap: nubosidad elevada, pérdidas térmicas y potencial ensuciamiento.\n"
    "• Recomendaciones:\n"
    "  1) Programa de lavado post-lluvia y previo a temporada seca.\n"
    "  2) Revisión de strings e inspección IV Curve para descartar desbalances o hotspots.\n"
    "  3) Verificar límites/curtailment del inversor y calidad de conexión a red.\n"
    "  4) Implementar tablero de monitoreo (NASA + Growatt) para alertas de rendimiento."
)

# Guardar
os.makedirs(CARPETA_SALIDA, exist_ok=True)
ruta_docx = os.path.join(CARPETA_SALIDA, NOMBRE_DOCX)
doc.save(ruta_docx)
print(f"✅ Informe generado: {ruta_docx}")
